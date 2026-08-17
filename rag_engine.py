
import os
import re
import hashlib
from datetime import datetime, timezone

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

ALLOWED_EXTENSIONS = {"pdf", "docx", "txt", "md"}
MAX_PROCESS_BYTES = 50 * 1024 * 1024
CHUNK_SIZE = 1200
CHUNK_OVERLAP = 180


def _get_supabase():
    from supabase import create_client
    url = os.getenv("SUPABASE_URL", "").strip()
    key = os.getenv("SUPABASE_SECRET_KEY", "").strip()
    if not url or not key:
        raise RuntimeError("SUPABASE_URL dan SUPABASE_SECRET_KEY belum tersedia.")
    return create_client(url, key)


def _clean_text(text):
    text = text or ""
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _validate_filename(filename):
    filename = filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError("Format yang didukung: PDF, DOCX, TXT, dan MD.")
    return ext


def extract_text_bytes(filename, raw):
    ext = _validate_filename(filename)
    if not raw:
        raise ValueError("Dokumen kosong.")
    if len(raw) > MAX_PROCESS_BYTES:
        raise ValueError("Ukuran dokumen melebihi batas pemrosesan 50 MB.")

    if ext == "pdf":
        try:
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(raw))
            text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
        except Exception as exc:
            raise RuntimeError(f"PDF tidak dapat dibaca: {exc}") from exc

    elif ext == "docx":
        try:
            from docx import Document
            import io
            doc = Document(io.BytesIO(raw))
            parts = [p.text for p in doc.paragraphs if p.text]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            text = "\n".join(parts)
        except Exception as exc:
            raise RuntimeError(f"DOCX tidak dapat dibaca: {exc}") from exc

    else:
        text = raw.decode("utf-8", errors="ignore")

    text = _clean_text(text)
    if len(text) < 30:
        raise ValueError(
            "Dokumen berhasil diunggah tetapi teks yang dapat diekstrak terlalu sedikit. "
            "Jika PDF berupa scan/gambar, diperlukan OCR."
        )
    return text, hashlib.sha256(raw).hexdigest()


def chunk_text(text, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    words = text.split()
    chunks, current, current_len = [], [], 0

    for word in words:
        current.append(word)
        current_len += len(word) + 1
        if current_len >= size:
            chunk = " ".join(current).strip()
            if chunk:
                chunks.append(chunk)

            keep, keep_len = [], 0
            for w in reversed(current):
                if keep_len >= overlap:
                    break
                keep.insert(0, w)
                keep_len += len(w) + 1
            current = keep
            current_len = keep_len

    if current:
        chunks.append(" ".join(current).strip())
    return chunks


def save_document(filename, text, sha256, uploaded_by, storage_path="", file_size=0):
    supabase = _get_supabase()
    chunks = chunk_text(text)
    now = datetime.now(timezone.utc).isoformat()

    existing = (
        supabase.table("ai_documents")
        .select("id")
        .eq("storage_path", storage_path)
        .limit(1)
        .execute()
    )
    if getattr(existing, "data", None):
        raise ValueError("Dokumen dengan path tersebut sudah ada di Knowledge Base.")

    doc_row = {
        "title": filename[:200],
        "filename": filename[:200],
        "storage_path": storage_path[:500],
        "file_type": filename.rsplit(".", 1)[-1].lower() if "." in filename else "",
        "file_size": int(file_size or 0),
        "uploaded_by": (uploaded_by or "-")[:200],
        "uploaded_at": now,
        "status": "processing",
        "chunk_count": 0,
        "error_message": None,
    }
    inserted = supabase.table("ai_documents").insert(doc_row).execute()
    if not getattr(inserted, "data", None):
        raise RuntimeError("Gagal membuat metadata dokumen di Supabase.")

    doc = inserted.data[0]
    doc_id = doc["id"]

    try:
        rows = [
            {
                "document_id": doc_id,
                "chunk_index": i,
                "content": chunk,
                "source": filename[:200],
            }
            for i, chunk in enumerate(chunks)
        ]

        # Insert bertahap agar payload database tetap kecil.
        for start in range(0, len(rows), 100):
            supabase.table("ai_chunks").insert(rows[start:start + 100]).execute()

        supabase.table("ai_documents").update({
            "status": "ready",
            "chunk_count": len(chunks),
            "error_message": None,
        }).eq("id", doc_id).execute()

    except Exception as exc:
        supabase.table("ai_documents").update({
            "status": "error",
            "error_message": str(exc)[:1000],
        }).eq("id", doc_id).execute()
        raise

    return doc_id, len(chunks)


def save_manual_knowledge(title, text, created_by):
    text = _clean_text(text)
    title = (title or "").strip()
    if len(title) < 3 or len(title) > 200:
        raise ValueError("Judul informasi tidak valid.")
    if len(text) < 30 or len(text) > 100000:
        raise ValueError("Isi informasi harus 30-100.000 karakter.")

    sha256 = hashlib.sha256(text.encode("utf-8")).hexdigest()
    return save_document(
        title,
        text,
        sha256,
        created_by,
        storage_path=f"manual/{sha256}.txt",
        file_size=len(text.encode("utf-8")),
    )


def list_documents():
    supabase = _get_supabase()
    result = (
        supabase.table("ai_documents")
        .select("*")
        .order("uploaded_at", desc=True)
        .execute()
    )
    return getattr(result, "data", None) or []


def delete_document(doc_id):
    supabase = _get_supabase()
    result = (
        supabase.table("ai_documents")
        .select("storage_path")
        .eq("id", doc_id)
        .limit(1)
        .execute()
    )
    rows = getattr(result, "data", None) or []
    if not rows:
        raise ValueError("Dokumen tidak ditemukan.")

    storage_path = rows[0].get("storage_path") or ""
    supabase.table("ai_documents").delete().eq("id", doc_id).execute()

    if storage_path and not storage_path.startswith("manual/"):
        try:
            supabase.storage.from_("ai-guides").remove([storage_path])
        except Exception as exc:
            print(f"[RAG] Metadata terhapus tetapi file Storage gagal dihapus: {exc}")

    return True


def retrieve(query, top_k=6, min_score=0.10):
    """
    Retrieval RAG tahap ini menggunakan TF-IDF atas chunks yang tersimpan
    di Supabase PostgreSQL. pgvector sudah dapat diaktifkan untuk tahap
    semantic embeddings berikutnya tanpa mengubah skema dokumen/chunks.
    """
    supabase = _get_supabase()
    result = (
        supabase.table("ai_chunks")
        .select("id,document_id,chunk_index,content,source")
        .limit(5000)
        .execute()
    )
    rows = getattr(result, "data", None) or []
    if not rows:
        return []

    corpus = [str(x.get("content") or "") for x in rows]
    if not any(corpus):
        return []

    try:
        vectorizer = TfidfVectorizer(
            lowercase=True,
            ngram_range=(1, 2),
            max_features=20000,
            sublinear_tf=True,
        )
        matrix = vectorizer.fit_transform(corpus + [query])
        scores = cosine_similarity(matrix[-1], matrix[:-1]).ravel()
    except Exception:
        q = set(re.findall(r"\w+", query.lower()))
        scores = []
        for text in corpus:
            t = set(re.findall(r"\w+", text.lower()))
            scores.append(len(q & t) / max(1, len(q)))

    ranked = sorted(
        zip(scores, rows),
        key=lambda x: x[0],
        reverse=True,
    )

    return [
        {
            "doc_id": item.get("document_id"),
            "chunk_id": item.get("id"),
            "text": item.get("content", ""),
            "source": item.get("source", "Dokumen panduan"),
            "score": round(float(score), 4),
        }
        for score, item in ranked[:top_k]
        if float(score) >= min_score
    ]


def build_context(results):
    if not results:
        return ""
    blocks = []
    for i, item in enumerate(results, 1):
        blocks.append(
            f"[Sumber {i}: {item['source']} | relevansi {item['score']}]\n"
            f"{item['text']}"
        )
    return "\n\n".join(blocks)
